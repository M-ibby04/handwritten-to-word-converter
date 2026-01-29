import gradio as gr
import os
import base64
from io import BytesIO
from groq import Groq
from docx import Document
from docx.shared import Pt

# Initialize Groq Client
client = Groq(api_key=os.getenv("GROQ_API_KEY"))

def encode_image(image):
    buffered = BytesIO()
    image.save(buffered, format="JPEG")
    return base64.b64encode(buffered.getvalue()).decode('utf-8')

def convert_handwriting(image):
    if not os.getenv("GROQ_API_KEY"):
        return "Error: Please add GROQ_API_KEY to Secrets.", None

    base64_image = encode_image(image)
    
    try:
        # Using the Llama 4 Scout model you verified
        model_id = "meta-llama/llama-4-scout-17b-16e-instruct" 
        
        prompt = """
        You are a specialized OCR assistant for FAST NUCES chemistry labs.
        Transcribe the attached handwritten notes:
        1. Maintain the two-column structure (read left then right).
        2. Format all chemical formulas and math in LaTeX (e.g., ΔH, Tc).
        3. Use Markdown style headers (# Title, ## Section).
        4. Identify if there are any diagrams and label them as [Diagram: Description].
        """
        
        completion = client.chat.completions.create(
            model=model_id,
            messages=[{"role": "user", "content": [
                {"type": "text", "text": prompt},
                {"type": "image_url", "image_url": {"url": f"data:image/jpeg;base64,{base64_image}"}}
            ]}],
            temperature=0.1, 
        )
        
        extracted_text = completion.choices[0].message.content
        
        # Word Document Generation (IEEE-ish formatting)
        doc = Document()
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(10) # IEEE Standard

        doc.add_heading('Chemistry Note Transcription', 0)
        for line in extracted_text.split('\n'):
            if line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), level=1)
            else:
                doc.add_paragraph(line)
            
        file_path = "Handwritten_Notes_IEEE.docx"
        doc.save(file_path)
        return extracted_text, file_path

    except Exception as e:
        return f"Error: {str(e)}", None

# Simple GUI
demo = gr.Interface(
    fn=convert_handwriting,
    inputs=gr.Image(type="pil", label="Upload Notes"),
    outputs=[gr.Textbox(label="Preview Text", lines=15), gr.File(label="Download .docx")],
    title="Handwriting to Word Docx Converter",
    description="Turn any hand-written notes to Word Document even containing complex layouts and formulas"
)

if __name__ == "__main__":
    demo.launch()
